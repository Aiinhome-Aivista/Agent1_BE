"""
Document loader & chunker for runbook ingestion.

Supports the four formats accepted by the Create-Runbook UI:
    .pdf   - via pypdf
    .docx  - via python-docx
    .md    - plain UTF-8
    .txt   - plain UTF-8

Chunking strategy is character-window with overlap (RUNBOOK_CHUNK_CHARS,
RUNBOOK_CHUNK_OVERLAP from settings). This is intentionally simple –
runbooks are short, well-structured docs and a sliding window is
plenty for retrieval quality.
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Iterable

from app.core.config import settings

logger = logging.getLogger(__name__)

SUPPORTED_EXTS = {".pdf", ".docx", ".md", ".txt"}


class DocumentService:
    # ------------------------------------------------------------------
    # Extraction
    # ------------------------------------------------------------------

    def extract_text(self, file_path: str | Path) -> str:
        """Dispatch to the right loader based on extension."""
        p = Path(file_path)
        ext = p.suffix.lower()

        if ext == ".pdf":
            return self._extract_pdf(p)
        if ext == ".docx":
            return self._extract_docx(p)
        if ext in {".md", ".txt"}:
            return self._extract_text_file(p)

        raise ValueError(
            f"Unsupported file type {ext!r}. Allowed: {sorted(SUPPORTED_EXTS)}"
        )

    @staticmethod
    def _extract_pdf(path: Path) -> str:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        chunks: list[str] = []
        for page in reader.pages:
            try:
                t = page.extract_text() or ""
            except Exception as e:
                logger.warning("PDF page extraction failed on %s: %s", path.name, e)
                t = ""
            if t.strip():
                chunks.append(t)
        return "\n\n".join(chunks)

    @staticmethod
    def _extract_docx(path: Path) -> str:
        from docx import Document  # python-docx

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]

        # Pull table cell text too – many runbook authors love tables
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n".join(paragraphs)

    @staticmethod
    def _extract_text_file(path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="replace")

    # ------------------------------------------------------------------
    # Chunking
    # ------------------------------------------------------------------

    def chunk(
        self,
        text: str,
        chunk_chars: int | None = None,
        overlap: int | None = None,
    ) -> list[str]:
        """
        Character-window chunker with overlap. Tries to break on a paragraph
        boundary if one exists inside the last 200 chars of the window;
        otherwise falls back to a hard cut.
        """
        chunk_chars = chunk_chars or settings.RUNBOOK_CHUNK_CHARS
        overlap     = overlap     or settings.RUNBOOK_CHUNK_OVERLAP

        text = self._normalise(text)
        if not text:
            return []

        out: list[str] = []
        i = 0
        n = len(text)

        while i < n:
            end = min(i + chunk_chars, n)
            window = text[i:end]

            # If we're not at EOF, try to break on the nearest paragraph break
            if end < n:
                tail = window[-200:]
                break_at = tail.rfind("\n\n")
                if break_at != -1:
                    end = i + (len(window) - 200) + break_at
                    window = text[i:end]

            chunk = window.strip()
            if chunk:
                out.append(chunk)

            if end >= n:
                break

            i = max(end - overlap, i + 1)  # always make forward progress

        return out

    @staticmethod
    def _normalise(text: str) -> str:
        # Collapse runs of blank lines + trim trailing whitespace per line
        text = re.sub(r"[ \t]+\n", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


document_service = DocumentService()
